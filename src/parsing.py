"""Document parsing utilities for resume and job-description text extraction.

Supports PDF (via PyPDF2) and DOCX (via python-docx) files. Also handles
raw text pasted by the user. Every public function is pure and returns a
string so the rest of the pipeline can stay format-agnostic.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

import streamlit as st

import PyPDF2
from docx import Document

# ---------------------------------------------------------------------------
# Type alias for accepted inputs
# ---------------------------------------------------------------------------
FileLike = Union[str, Path, bytes, io.BytesIO, "st.runtime.uploaded_file_manager.UploadedFile"]


def _read_bytes(source: FileLike) -> bytes:
    """Normalise a variety of inputs into raw bytes.

    Handles ``pathlib.Path`` / ``str`` paths, in-memory bytes, ``BytesIO``
    buffers and Streamlit ``UploadedFile`` objects transparently.
    """
    if isinstance(source, (bytes, bytearray)):
        return bytes(source)

    if isinstance(source, io.BytesIO):
        return source.getvalue()

    if isinstance(source, (str, Path)):
        return Path(source).read_bytes()

    # Streamlit UploadedFile exposes a getvalue() method and a read() method.
    getvalue = getattr(source, "getvalue", None)
    if callable(getvalue):
        return getvalue()

    read = getattr(source, "read", None)
    if callable(read):
        return read()

    raise TypeError(f"Unsupported source type for byte extraction: {type(source)!r}")


def extract_text_from_pdf(source: FileLike) -> str:
    """Return concatenated text from every page of a PDF document."""
    raw = _read_bytes(source)
    reader = PyPDF2.PdfReader(io.BytesIO(raw))

    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n".join(pages)


def extract_text_from_docx(source: FileLike) -> str:
    """Return concatenated paragraph text from a DOCX document."""
    raw = _read_bytes(source)
    document = Document(io.BytesIO(raw))

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Tables often hold experience / skills in ATS-style resumes.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    return "\n".join(paragraphs)


def extract_text(source: FileLike, filename: str | None = None) -> str:
    """Auto-detect file type from ``filename`` and extract text.

    Falls back to PDF parsing when the type cannot be inferred. Raises
    ``ValueError`` for unsupported extensions.
    """
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(source)
    if name.endswith(".docx"):
        return extract_text_from_docx(source)

    # No filename hint: try PDF first (binary magic), then DOCX (zip magic).
    raw = _read_bytes(source)
    if raw.startswith(b"%PDF"):
        return extract_text_from_pdf(raw)
    if raw[:2] == b"PK":
        return extract_text_from_docx(raw)

    raise ValueError(
        "Unsupported file type. Please upload a .pdf or .docx file, "
        "or paste the text directly."
    )


def clean_extracted_text(text: str) -> str:
    """Collapse excessive whitespace and strip stray control characters."""
    if not text:
        return ""

    # Replace common PDF artefacts: hyphenation across line breaks.
    cleaned = text.replace("-\n", "").replace("\n", " ")
    # Collapse runs of whitespace.
    cleaned = " ".join(cleaned.split())
    return cleaned.strip()
